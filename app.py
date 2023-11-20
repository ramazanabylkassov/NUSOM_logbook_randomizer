import streamlit as st
import pandas as pd
import numpy as np
from datetime import timedelta
import io
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from streamlit_gsheets import GSheetsConnection

def main():
    hospitals = sorted(["Cardiac Research and Surgery Center", "UMC National Center for Maternal and Child Health", "Municipal Children's Hospital 2", "UMC Children's Rehabilitation Center", "Municipal Children's Hospital 3"])
    departments = { 'Pediatric Cardiology': {'tutor': 'Ivanova-Razumova T.V.', 'hospital': "Cardiac Research and Surgery Center"}, 
                    'Pediatric Gastroenterology': {'tutor': 'Ibrayeva A.K.', 'hospital': "UMC National Center for Maternal and Child Health"}, 
                    'Pediatric Rheumatology': {'tutor': 'Assylbekova M.K.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    'Pediatric Endocrinology': {'tutor': 'Rakhimzhanova M.K.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    'Pediatric Radiology': {'tutor': 'Dautov T.B.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    'Pediatric Emergencies': {'tutor': 'Baigiriyev R.M.', 'hospital': "Municipal Children's Hospital 2"},
                    'General Pediatrics and Pediatric Rehabilitation': {'tutor': 'Daribayev Zh.R.', 'hospital': "UMC Children's Rehabilitation Center"},
                    'Pediatric Infectious Diseases': {'tutor': 'Utegenova R.B.', 'hospital': "Municipal Children's Hospital 3"},
                    'Pediatric Oncology': {'tutor': 'Shaikhyzada K.K.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    'Pediatric Hematology': {'tutor': 'Umirbekova B.B.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    'Pediatric Nephrology': {'tutor': 'Rakhimzhanova S.S.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    'Pediatric Intensive Care Unit': {'tutor': 'Saparov A.I.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    'Neonatal Intensive Care Unit': {'tutor': 'Abentayeva B.A.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    'Pediatric Allergology, Immunology and Pulmonology': {'tutor': 'Kovzel E.F.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    'Neonatology': {'tutor': 'Tortayeva G.S.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    'Pediatric Neurology': {'tutor': 'Nauryzbayeva A.A.', 'hospital': "UMC National Center for Maternal and Child Health"},
                    }

    with st.sidebar:
        name = st.text_input('Enter your first & last name:', value='Dimitri Poddighe')
        department = st.selectbox('Choose department:', sorted(departments.keys(), key=str.lower))
        tutor = st.text_input("Enter clinical preceptor's name:", value=departments[department]['tutor'])
        hospital = st.selectbox("Choose hospital:", hospitals, index=hospitals.index(departments[department]['hospital']))
        cols = st.columns(2)
        start_date = cols[0].date_input('Enter start date')
        end_date = cols[1].date_input('Enter end date', value=start_date+timedelta(days=54))
        
        if department == 'Pediatric Radiology':
            patient_amount_us = st.slider("Ultrasound:", min_value=1, max_value=100, step=1, value=30)
            patient_amount_xray = st.slider("X-ray:", min_value=1, max_value=100, step=1, value=30)
            patient_amount_CT = st.slider("CT:", min_value=1, max_value=100, step=1, value=30)
            patient_amount_MRI = st.slider("MRI:", min_value=1, max_value=100, step=1, value=30)
            patient_amount = patient_amount_us + patient_amount_xray + patient_amount_CT + patient_amount_MRI
        else:    
            patient_amount = st.slider("Select the amount of patients:", min_value=1, max_value=100, step=1, value=30)
        
        if department in ['Neonatal Intensive Care Unit', 'Neonatology']:
            age_range = np.arange(0, 28, 1)
        else:
            age_range = np.arange(0.1, 17, 0.1)
        if st.toggle("Preferentiate certain age group"):
            if department in ['Neonatal Intensive Care Unit', 'Neonatology']:
                st.write(f'This option is not available for the {department} department')
            else:
                preferent_age_group = st.select_slider('Preferentiable age range (years)', options=np.arange(0, 18, step=0.5), value=(0, 17.5))
                preferent_probab = st.select_slider('Percent of patients in the selected age group (%)', options=np.arange(0, 101, step=1), value=50)
                remaining_step = (((preferent_age_group[0]-0.1) + (17.9 - preferent_age_group[1]))) / (patient_amount*(100-preferent_probab)/100)
                age_range = np.concatenate([np.linspace(preferent_age_group[0], preferent_age_group[1], num=round(patient_amount*preferent_probab/100)), np.arange(0.1, preferent_age_group[0], step=remaining_step), np.arange(preferent_age_group[1], 17.9, step=remaining_step)])

    sheet_name = department
    if sheet_name == 'General Pediatrics and Pediatric Rehabilitation':
        sheet_name = "Rehabilitation"
    elif sheet_name == 'Pediatric Allergology, Immunology and Pulmonology':
        sheet_name = "Allergology"
    elif sheet_name == 'Pediatric Radiology':
        sheet_name_us = 'radiology_US'
        sheet_name_xray = 'radiology_X_ray'
        sheet_name_ct = 'radiology_CT'
        sheet_name_mri = 'radiology_MRI'

    # # Create a connection object.
    # url = "https://docs.google.com/spreadsheets/d/1m0JxgnEUoojl_o1eRC7zRyJxF-7SxkfTS6qLUVq8Poc/edit?usp=sharing"
    # conn = st.connection("gsheets", type=GSheetsConnection)
    # df_diseases = conn.read(spreadsheet=url, usecols=[0, 1])

    if sheet_name == 'Pediatric Radiology':
        df_diseases_us = pd.read_excel('logbook_common_diseases.xlsx', sheet_name=sheet_name_us)
        df_diseases_xray = pd.read_excel('logbook_common_diseases.xlsx', sheet_name=sheet_name_xray)
        df_diseases_ct = pd.read_excel('logbook_common_diseases.xlsx', sheet_name=sheet_name_ct)
        df_diseases_mri = pd.read_excel('logbook_common_diseases.xlsx', sheet_name=sheet_name_mri)
    else:
        df_diseases = pd.read_excel('logbook_common_diseases.xlsx', sheet_name=sheet_name)
    
    output = {
                "№": None,
                "Age": None,
                "Gender": None,
                "Diagnosis": None,
                "Tutor's name": None,
                }

    patient_index = np.arange(1, patient_amount+1, 1)
    patients_age_numeric = np.random.choice(age_range, size=patient_amount)
    patients_age = np.empty_like(patients_age_numeric, dtype='object')
    for index, age in enumerate(patients_age_numeric):
        if department in ['Neonatology', 'Neonatal Intensive Care Unit']:
            patients_age[index] = f'{round(age)} days'
        elif age >= 1:
            patients_age[index] = f'{round(age)} {"years" if age > 1 else "year"}'
        else:
            patients_age[index] = f'{round(age*12)} months'

    patients_gender = np.random.choice(['male', 'female'], size=patient_amount)
    
    if sheet_name == 'Pediatric Radiology':
        diseases = np.concatenate((
            np.random.choice(df_diseases_us['Common diseases'], size=patient_amount_us, p=df_diseases_us['Frequency']),
            np.random.choice(df_diseases_xray['Common diseases'], size=patient_amount_xray, p=df_diseases_xray['Frequency']),
            np.random.choice(df_diseases_ct['Common diseases'], size=patient_amount_CT, p=df_diseases_ct['Frequency']),
            np.random.choice(df_diseases_mri['Common diseases'], size=patient_amount_MRI, p=df_diseases_mri['Frequency'])           
        ))
    else:
        diseases = np.random.choice(df_diseases['Common diseases'], size=patient_amount, p=df_diseases['Frequency'])

    if sheet_name == 'Pediatric Radiology':
        tutor_list = np.concatenate((
            ['Yurana Albayeva']*patient_amount_us,
            ['Shaddat Umbetov']*patient_amount_xray,
            ['Baurzhan Kaliev']*patient_amount_CT,
            [tutor]*patient_amount_MRI,
        ))
    else:
        tutor_list = [tutor]*patient_amount
    

    output['№'] = patient_index
    output['Age'] = patients_age
    output['Gender'] = patients_gender
    output['Diagnosis'] = diseases
    output["Tutor's name"] = tutor_list

### ------- MAIN PAGE -----------
    header = st.container()
    
    header_text = (f'''
             ***<h3 style="text-align: center;"> LOGBOOK</h3>***
             ***RESIDENT:*** {name} \n
             ***ROTATION:*** {department} \n
             ***HOSPITAL SITE:*** {hospital} \n
             ***Attendance Date:*** **From:** {start_date}  **To:** {end_date} \n
             ***Supervisor (name and surname):*** {tutor} \n
             ***Supervisor (signature):*** \n
             ###
             ''')

    header.markdown(header_text, unsafe_allow_html=True)
    
    df_output = pd.DataFrame(output)

    st.data_editor(df_output, use_container_width=True)

    # Initialise the Word document
    doc = docx.Document()

    header = doc.add_paragraph("LOGBOOK")
    header.runs[0].bold = True  # Set the first run (text) to bold
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    resident = doc.add_paragraph()
    resident.add_run('RESIDENT: ').bold = True
    resident.add_run(f'{name}')
    rotation = doc.add_paragraph()
    rotation.add_run('ROTATION: ').bold = True
    rotation.add_run(f'{department}')
    hospital_site = doc.add_paragraph()
    hospital_site.add_run('HOSPITAL SITE: ').bold = True
    hospital_site.add_run(f'{hospital}')
    attendance = doc.add_paragraph()
    attendance.add_run('Attendance Date: ').bold = True
    attendance.add_run('From: ').bold = True
    attendance.add_run(f'{start_date} ')
    attendance.add_run('To: ').bold = True
    attendance.add_run(f'{end_date}')
    supervisor = doc.add_paragraph()
    supervisor.add_run('Supervisor (name and surname): ').bold = True
    supervisor.add_run(f'{tutor}')
    signature  = doc.add_paragraph()
    signature.add_run('Supervisor (signature): ').bold = True

    # Initialise the table
    t = doc.add_table(rows=(df_output.shape[0] + 1), cols=df_output.shape[1])
    # Add borders
    t.style = 'TableGrid'
    # Add the column headings
    for j in range(df_output.shape[1]):
        cell = t.cell(0, j)
        cell.text = df_output.columns[j]
        cell.paragraphs[0].runs[0].bold = True  # Set the first run (text) in the cell to bold

    # Add the body of the data frame
    for i in range(df_output.shape[0]):
        for j in range(df_output.shape[1]):
            cell = df_output.iat[i, j]
            t.cell(i + 1, j).text = str(cell)

    for cell in t.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in t.columns[2].cells:
        cell.width = Inches(1)
    for cell in t.columns[3].cells:
        cell.width = Inches(2)
    for cell in t.columns[4].cells:
        cell.width = Inches(3)

    # Center-align the content in the table
    for row in t.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    bio = io.BytesIO()
    doc.save(bio)
    if doc:
        st.download_button(
            label="Click here to download",
            data=bio.getvalue(),
            file_name=f"logbook_patients_{department}.docx",
            mime="docx"
        )

if __name__ == '__main__':
    main()